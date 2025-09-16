# mstautogen.py
import streamlit as st
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.chains import RetrievalQA
from langchain.vectorstores import FAISS
from langchain_community.embeddings import HuggingFaceEmbeddings
from langchain.llms import HuggingFacePipeline
from transformers import AutoTokenizer, AutoModelForSeq2SeqLM, pipeline
from PyPDF2 import PdfReader
from io import BytesIO
from docx import Document
from reportlab.platypus import SimpleDocTemplate, Paragraph
from reportlab.lib.styles import getSampleStyleSheet

# --------------------------------------
# Streamlit UI
# --------------------------------------
st.set_page_config(page_title="MST AutoGen (Amana)", layout="wide")

st.markdown(
    """
    <div style="background-color:#1E90FF;padding:12px;border-radius:8px;text-align:center;">
        <h1 style="color:white;">📑 METHOD STATEMENT AUTOGENERATION (AMANA AQAI)</h1>
    </div>
    """,
    unsafe_allow_html=True,
)

st.markdown("---")

# ----------------------------
# 1️⃣ Upload PDF(s)
# ----------------------------
uploaded_files = st.file_uploader(
    "📂 UPLOAD SPECIFICATIONS PDFs",
    accept_multiple_files=True,
    type=["pdf"],
    help="Upload contract documents in PDF format"
)

texts = []
if uploaded_files:
    for file in uploaded_files:
        pdf = PdfReader(file)
        text = ""
        for page in pdf.pages:
            page_text = page.extract_text()
            if page_text:
                text += page_text + "\n"
        texts.append(text)
    st.success(f"✅ Loaded {len(uploaded_files)} PDF(s).")

# ----------------------------
# 2️⃣ Split Text & Create Embeddings
# ----------------------------
if texts:
    st.info("⏳ Analysing the specifications and preparing Method Statement...")
    splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=50)
    docs = []
    for t in texts:
        docs.extend(splitter.split_text(t))

    embeddings = HuggingFaceEmbeddings(model_name="sentence-transformers/all-MiniLM-L6-v2")
    vectorstore = FAISS.from_texts(docs, embeddings)
    st.success("✅ Specifications analysed and Project specific Method statement created!")

# ----------------------------
# 3️⃣ Load Local LLM
# ----------------------------
@st.cache_resource
def load_local_llm():
    model_name = "google/flan-t5-small"  # you can upgrade to flan-t5-base/large
    tokenizer = AutoTokenizer.from_pretrained(model_name)
    model = AutoModelForSeq2SeqLM.from_pretrained(model_name)
    pipe = pipeline("text2text-generation", model=model, tokenizer=tokenizer, max_length=1024)
    return HuggingFacePipeline(pipeline=pipe)

if texts:
    llm = load_local_llm()

# ----------------------------
# 4️⃣ Generate Method Statement
# ----------------------------
if texts:
    project_title = st.text_input("📝 Activity Title for the Method Statement")

    if project_title:
        st.info("⏳ Generating Method Statement as per specifications and contract documents...")

        retriever = vectorstore.as_retriever()
        qa = RetrievalQA.from_chain_type(llm=llm, retriever=retriever)

        # Define structured prompt
        ms_prompt = f"""
        Based on the contract specifications provided, generate a professional Method Statement
        for the activity: "{project_title}".

        The Method Statement must include the following sections:

        1. Introduction
        2. Purpose
        3. Scope
        4. Tools
        5. Materials Used
        6. Execution Method
        7. Post Execution

        Ensure content is clear, concise, and aligned with typical construction/commissioning practices.
        """

        answer = qa.run(ms_prompt)

        st.subheader("📄 Generated Method Statement")
        st.write(answer)

        # Structured format
        structured_prompt = f"""
        Format the Method Statement below into a structured outline with proper headings:

        {answer}

        Format strictly as:

        ## Introduction
        ...
        ## Purpose
        ...
        ## Scope
        ...
        ## Tools
        ...
        ## Materials Used
        ...
        ## Execution Method
        ...
        ## Post Execution
        """

        structured_output = llm(structured_prompt)
        st.subheader("📋 Structured Method Statement")
        st.markdown(structured_output)

        # ----------------------------
        # 5️⃣ Export Options
        # ----------------------------

        def export_docx(content):
            buffer = BytesIO()
            doc = Document()
            doc.add_heading("Method Statement", 0)
            for line in content.split("\n"):
                doc.add_paragraph(line)
            doc.save(buffer)
            buffer.seek(0)
            return buffer

        def export_pdf(content):
            buffer = BytesIO()
            doc = SimpleDocTemplate(buffer)
            styles = getSampleStyleSheet()
            flowables = []
            for line in content.split("\n"):
                flowables.append(Paragraph(line, styles["Normal"]))
            doc.build(flowables)
            buffer.seek(0)
            return buffer

        col1, col2 = st.columns(2)
        with col1:
            docx_file = export_docx(structured_output)
            st.download_button(
                label="⬇️ Download as Word",
                data=docx_file,
                file_name=f"Method_Statement_{project_title}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
        with col2:
            pdf_file = export_pdf(structured_output)
            st.download_button(
                label="⬇️ Download as PDF",
                data=pdf_file,
                file_name=f"Method_Statement_{project_title}.pdf",
                mime="application/pdf"
            )

